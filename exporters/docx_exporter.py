"""DOCX export generation."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.models.schemas import BlockType, PageContent


class DocxExporter:
    """Create a readable, navigable Word document."""

    def export(self, document_title: str, pages: list[PageContent], include_metadata: bool) -> bytes:
        document = Document()
        self._configure_styles(document)
        self._add_cover(document, document_title, len(pages))
        self._add_summary(document, pages)
        self._add_toc(document)

        for index, page in enumerate(pages, start=1):
            document.add_page_break()
            document.add_heading(page.title, level=1)
            source_line = document.add_paragraph()
            source_line.style = document.styles["Intense Quote"]
            source_line.add_run("Source URL: ").bold = True
            source_line.add_run(page.final_url)

            if include_metadata and (page.meta_description or page.publication_date):
                metadata = document.add_paragraph()
                metadata.style = document.styles["Normal"]
                if page.publication_date:
                    metadata.add_run(f"Published: {page.publication_date}  ")
                if page.meta_description:
                    metadata.add_run(f"Summary: {page.meta_description}")

            for block in page.blocks:
                self._append_block(document, block)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _configure_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)

        for style_name, size in [("Title", 24), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
            style = document.styles[style_name]
            style.font.name = "Aptos"
            style.font.size = Pt(size)

        for section in document.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

    def _add_cover(self, document: Document, title: str, page_count: int) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(24)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Readable website export").italic = True

        count_line = document.add_paragraph()
        count_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        count_line.add_run(f"Pages captured: {page_count}")

    def _add_summary(self, document: Document, pages: list[PageContent]) -> None:
        document.add_paragraph()
        document.add_heading("Summary", level=1)
        if not pages:
            document.add_paragraph("No pages were successfully scraped.")
            return

        summary = document.add_paragraph(style="List Bullet")
        summary.add_run(f"Source: {pages[0].requested_url}")
        document.add_paragraph(f"Pages included: {len(pages)}", style="List Bullet")
        document.add_paragraph(
            f"Total words: {sum(page.word_count for page in pages):,}",
            style="List Bullet",
        )

    def _add_toc(self, document: Document) -> None:
        document.add_page_break()
        document.add_heading("Contents", level=1)
        paragraph = document.add_paragraph()

        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        hint = OxmlElement("w:t")
        hint.text = "Update fields in Word to populate the table of contents."
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(hint)
        run._r.append(fld_char_end)

    def _append_block(self, document: Document, block) -> None:
        if block.kind == BlockType.HEADING and block.text:
            level = 2 if (block.level or 2) <= 2 else 3
            document.add_heading(block.text, level=level)
        elif block.kind == BlockType.PARAGRAPH and block.text:
            document.add_paragraph(block.text)
        elif block.kind == BlockType.QUOTE and block.text:
            quote = document.add_paragraph(style="Intense Quote")
            quote.add_run(block.text)
        elif block.kind == BlockType.BULLET_LIST:
            for item in block.items:
                document.add_paragraph(item, style="List Bullet")
        elif block.kind == BlockType.ORDERED_LIST:
            for item in block.items:
                document.add_paragraph(item, style="List Number")
        elif block.kind == BlockType.TABLE and block.rows:
            table = document.add_table(rows=1, cols=len(block.rows[0]))
            table.style = "Table Grid"
            for col, value in enumerate(block.rows[0]):
                table.rows[0].cells[col].text = value
            for row in block.rows[1:]:
                cells = table.add_row().cells
                for col, value in enumerate(row):
                    cells[col].text = value

